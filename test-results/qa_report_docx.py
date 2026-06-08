from pathlib import Path
from docx import Document

p = Path(r"\\wsl$\Ubuntu\home\erra\www\Alghazaly\test-results\Alghazaly_DB_E2E_Test_Report.docx")
doc = Document(p)
print("exists", p.exists())
print("size", p.stat().st_size)
print("paragraphs", len(doc.paragraphs))
print("tables", len(doc.tables))
print("title", doc.paragraphs[0].text)
print(
    "summary_present",
    any("15/15 passed" in para.text for para in doc.paragraphs)
    or any("15/15 passed" in cell.text for table in doc.tables for row in table.rows for cell in row.cells),
)
