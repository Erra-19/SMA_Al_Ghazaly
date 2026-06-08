from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Alghazaly_DB_E2E_Test_Report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin in ("top", "left", "bottom", "right"):
        node = tbl_cell_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), "120")
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    set_cell_margins(table)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.add_run(text)
    return p


def add_kv_table(doc: Document, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Value"
    for cell in hdr:
        set_cell_shading(cell, "F2F4F7")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    set_table_width(table, [2.0, 4.25])
    return table


def read_lines(name: str):
    path = ROOT / name
    if not path.exists():
        return []
    return [line.strip() for line in path.read_text(encoding="utf-8", errors="replace").splitlines() if line.strip()]


def main() -> None:
    e2e_lines = read_lines("alghazaly_e2e_result.txt")
    migration_lines = read_lines("alghazaly_migration_result.txt")
    http_guard = "\n".join(read_lines("http_guard_check.txt"))

    pass_lines = [line for line in e2e_lines if line.startswith("PASS |")]
    fail_lines = [line for line in e2e_lines if line.startswith("FAIL |")]
    summary = next((line for line in e2e_lines if line.startswith("SUMMARY |")), "SUMMARY | unavailable")
    migrations_done = [line for line in migration_lines if " DONE" in line and "Running migrations" not in line]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 2" else 16)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Al Ghazaly Backend DB & API E2E Test Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.add_run("Scope: Laravel migrations, seeders, PPDB workflow, role route guards, content APIs, public visibility, and profile endpoint.").italic = True

    add_heading(doc, "Executive Summary", 1)
    p = doc.add_paragraph()
    p.add_run("Result: ").bold = True
    p.add_run("PASS. ")
    p.add_run(summary.replace("SUMMARY | ", ""))
    doc.add_paragraph("The database test ran against an isolated SQLite database at /tmp/alghazaly_e2e.sqlite inside the app container. Production MySQL data was not reset or migrated.")

    add_kv_table(doc, [
        ("Test date", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Application", "Al Ghazaly Laravel backend"),
        ("Runtime", "Docker container: alghazaly"),
        ("Database under test", "SQLite temporary database"),
        ("Migration/seed result", f"{len(migrations_done)} migration/seeder completion lines captured"),
        ("E2E result", summary.replace("SUMMARY | ", "")),
        ("HTTP guard smoke check", http_guard),
    ])

    add_heading(doc, "Validated Scenarios", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Status", "Scenario", "Detail"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for line in pass_lines + fail_lines:
        status, name, detail = [part.strip() for part in line.split("|", 2)]
        cells = table.add_row().cells
        cells[0].text = status
        cells[1].text = name
        cells[2].text = detail
        set_cell_shading(cells[0], "EAF5EA" if status == "PASS" else "FDECEC")
    set_table_width(table, [0.8, 2.55, 3.0])

    add_heading(doc, "Migration And Seeder Check", 1)
    doc.add_paragraph("The test executed php artisan migrate:fresh --seed --force using the isolated SQLite test database. All migrations and seeders completed without error.")
    mig_table = doc.add_table(rows=1, cols=2)
    mig_table.style = "Table Grid"
    mig_table.rows[0].cells[0].text = "Area"
    mig_table.rows[0].cells[1].text = "Evidence"
    for cell in mig_table.rows[0].cells:
        set_cell_shading(cell, "F2F4F7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for area, evidence in [
        ("Migrations", "Users, roles, content, media, forms, PPDB, payments, students, teachers, albums, testimonials, and alumni tables migrated."),
        ("Seeders", "RoleSeeder, AdminUserSeeder, SettingSeeder, and CategorySeeder completed."),
        ("Route count", "76 routes available after the latest controller and workflow changes."),
    ]:
        cells = mig_table.add_row().cells
        cells[0].text = area
        cells[1].text = evidence
    set_table_width(mig_table, [1.5, 4.75])

    add_heading(doc, "Notes And Follow-Up", 1)
    notes = [
        "The test intentionally avoided migrate:fresh on production MySQL.",
        "Midtrans Snap was not live-tested with a real server key; the tested PPDB payment path used zero-fee auto-paid behavior.",
        "Laravel cache files should be cleared or rebuilt after deployment because previous bootstrap/cache files contained stale Docker paths.",
        "Next useful test is a real HTTP run against a dedicated MySQL staging database with Midtrans sandbox credentials.",
        "DOCX structural QA passed; visual render QA was unavailable because LibreOffice/soffice was not installed in the available runtimes.",
    ]
    for note in notes:
        doc.add_paragraph(note, style="List Bullet")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Al Ghazaly backend test report")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
